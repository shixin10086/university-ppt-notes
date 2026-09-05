#!/usr/bin/env python3
"""Export cumulative Markdown notes as a readable Word document."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PAGE_RE = re.compile(r"^## P(\d+)\s+(.+?)\s*$", re.M)
LABEL_RE = re.compile(r"^【([^】]+)】\s*$", re.M)


def set_font(run, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def remove_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for name in ("top", "left", "bottom", "right", "between"):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "nil")
        borders.append(edge)
    p_pr.append(borders)


def parse(path: Path) -> tuple[str, list[dict]]:
    raw = path.read_text(encoding="utf-8")
    title_match = re.search(r"^#\s+(.+?)\s*$", raw, re.M)
    title = title_match.group(1) if title_match else "逐页授课备稿"
    pages = []
    matches = list(PAGE_RE.finditer(raw))
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(raw)
        section = raw[match.end():end]
        blocks = []
        labels = list(LABEL_RE.finditer(section))
        for label_index, label in enumerate(labels):
            block_end = labels[label_index + 1].start() if label_index + 1 < len(labels) else len(section)
            text = re.sub(r"\n\s*---\s*$", "", section[label.end():block_end].strip()).strip()
            if text:
                blocks.append((label.group(1), text))
        pages.append((int(match.group(1)), match.group(2).strip(), blocks))
    return title, pages


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    title, pages = parse(args.input)
    if not pages:
        raise SystemExit("No slide headings found.")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(7)

    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(130)
    remove_border(title_p)
    set_font(title_p.add_run(title), 23, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(subtitle.add_run(f"共 {len(pages)} 页"), 10.5, color="666666")
    doc.add_page_break()

    for slide, page_title, blocks in pages:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(5)
        heading.paragraph_format.keep_with_next = True
        set_font(heading.add_run(f"P{slide:02d}  {page_title}"), 15, bold=True)
        for label, text in blocks:
            label_p = doc.add_paragraph()
            label_p.paragraph_format.keep_with_next = True
            label_p.paragraph_format.space_after = Pt(2)
            set_font(label_p.add_run(f"【{label}】"), 10.5, bold=True, color="1F4E79")
            for line in filter(None, (part.strip() for part in text.splitlines())):
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.first_line_indent = Cm(0.74)
                paragraph.paragraph_format.keep_together = True
                set_font(paragraph.add_run(line), 11)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = title
    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
